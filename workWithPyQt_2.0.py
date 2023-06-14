from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import *
from PyQt5 import uic
import sys
from urllib.request import urlretrieve
import time
import sqlite3
import docx
import os
import openpyxl
from PyQt5.QtCore import *

choses_file = []

file_name = []


class MyThread(QtCore.QThread):
    mysignal = QtCore.pyqtSignal(str)
    def  __init__(self, parent=None):
        QtCore.QThread.__init__(self, parent)
    def run(self):
      try:
          print(file_name[0])
          db = sqlite3.connect('FIS.db')
          cur = db.cursor()
          if file_name[0].endswith(".docx"):
            doc = docx.Document(file_name[0])
            text = []
            for table in doc.tables:
              for row in table.rows:
                x = []
                if len(x) < 5:
                  for cell in row.cells:
                    x.append(cell.text)
                  else:  
                    cur.execute("INSERT INTO students (fam, imy, otch, ball, specialnost, forma) VALUES (?, ?, ?, ?, ?, ?)", (x[0], x[1], x[2], x[3], x[4], x[5]))
                    print(x)                     
                    db.commit() 
            db.close()


          elif file_name[0].endswith(".xlsx"):
            print("normis 2")
            db = sqlite3.connect('FIS.db')
            cur = db.cursor()
            file_to_read = openpyxl.load_workbook(file_name[0], data_only=True)
            sheet = file_to_read['Лист1']
            if str(file_to_read.sheetnames) == "['Лист1']":
              for row in range(2, sheet.max_row + 1):
                data = []
                for col in range(1, 7):
                    value = sheet.cell(row, col).value
                    data.append(value)
                ball = str(data[4])
                if len(ball) < 2:
                      ball = ball + '.0'
                cur.execute("INSERT INTO students (fam, imy, otch, ball, specialnost, forma) VALUES (?, ?, ?, ?, ?, ?)", (data[0], data[1], data[2], data[3], ball, data[5]))
                print(data)
                db.commit()
          else:
            print("Некорректное название листа")
      except IndexError:
          print('1')


class Ui_MainWindow(object): # Наследованный класс от объекта для отрисовк интерфейса

    def setupUi(self, MainWindow): # первоначальная функция с созданием интерфейса программы
        MainWindow.setObjectName("MainWindow")
        MainWindow.setFixedSize(1141, 221)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setStyleSheet("QPushButton#buttonMoveUp{\n"
"background-image: url(:/style/images/upfocusnormal.png);\n"
"background-position: center;\n"
"border: none;\n"
"width: 30px;\n"
"height: 30px;\n"
"}\n"
"\n"
"QPushButton#buttonMoveUp::hover {\n"
"background-image:url(:/style/images/upfocuspressed.png);\n"
"}\n"
"\n"
"QPushButton#buttonMoveUp::pressed {\n"
"background-image:url(:/style/images/upfocuspressed.png);\n"
"}\n"
"\n"
"QPushButton#buttonMoveUp::!enabled {\n"
"background-image:url(:/style/images/upnormal.png);\n"
"}\n"
"")
        self.centralwidget.setObjectName("centralwidget")
        self.tabWidget = QtWidgets.QTabWidget(self.centralwidget)
        self.tabWidget.setGeometry(QtCore.QRect(0, 0, 1171, 661))
        self.tabWidget.setMinimumSize(QtCore.QSize(1171, 0))
        self.tabWidget.setStyleSheet("background-color: qlineargradient(spread:pad, x1:1, y1:1, x2:1, y2:0, stop:0.00568182 rgba(255, 255, 255, 255), stop:1 rgba(151, 180, 255, 255));")
        self.tabWidget.setObjectName("tabWidget")
        self.tab = QtWidgets.QWidget()
        self.tab.setObjectName("tab")
        self.pushButton = QtWidgets.QPushButton(self.tab)
        self.pushButton.setGeometry(QtCore.QRect(20, 80, 261, 81))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.pushButton.setFont(font)
        self.pushButton.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(151, 180, 255, 255);")
        self.pushButton.setObjectName("pushButton")
        self.label = QtWidgets.QLabel(self.tab)
        self.label.setGeometry(QtCore.QRect(-10, -30, 1181, 751))
        self.label.setStyleSheet("background-color: qlineargradient(spread:pad, x1:1, y1:1, x2:1, y2:0, stop:0.00568182 rgba(255, 255, 255, 255), stop:1 rgba(151, 180, 255, 255));")
        self.label.setText("")
        self.label.setObjectName("label")
        self.textBrowser = QtWidgets.QTextBrowser(self.tab)
        self.textBrowser.setGeometry(QtCore.QRect(20, 10, 1101, 61))
        self.textBrowser.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: grey;\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.textBrowser.setObjectName("textBrowser")
        self.pushButton_2 = QtWidgets.QPushButton(self.tab)
        self.pushButton_2.setGeometry(QtCore.QRect(300, 80, 261, 81))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.pushButton_2.setFont(font)
        self.pushButton_2.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(151, 180, 255, 255);")
        self.pushButton_2.setObjectName("pushButton_2")
        self.zxc = QtWidgets.QPushButton(self.tab)
        self.zxc.setGeometry(QtCore.QRect(580, 80, 261, 81))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.zxc.setFont(font)
        self.zxc.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(151, 180, 255, 255);")
        self.zxc.setObjectName("zxc")
        self.pushButton_3 = QtWidgets.QPushButton(self.tab)
        self.pushButton_3.setGeometry(QtCore.QRect(860, 80, 261, 81))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.pushButton_3.setFont(font)
        self.pushButton_3.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(151, 180, 255, 255);")
        self.pushButton_3.setObjectName("pushButton_3")
        self.label_9 = QtWidgets.QLabel(self.tab)
        self.label_9.setGeometry(QtCore.QRect(0, 440, 271, 271))
        self.label_9.setText("")
        self.label_9.setPixmap(QtGui.QPixmap("81af2737-37a7-4a3a-a0b5-009923086361.png"))
        self.label_9.setScaledContents(True)
        self.label_9.setObjectName("label_9")
        self.label.raise_()
        self.pushButton.raise_()
        self.textBrowser.raise_()
        self.pushButton_2.raise_()
        self.zxc.raise_()
        self.pushButton_3.raise_()
        self.label_9.raise_()
        self.tabWidget.addTab(self.tab, "")
        self.tab_2 = QtWidgets.QWidget()
        self.tab_2.setObjectName("tab_2")
        self.pushButton_5 = QtWidgets.QPushButton(self.tab_2)
        self.pushButton_5.setGeometry(QtCore.QRect(550, 90, 561, 61))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.pushButton_5.setFont(font)
        self.pushButton_5.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"color: rgb(151, 180, 255, 255);")
        self.pushButton_5.setObjectName("pushButton_5")
        self.textEdit = QtWidgets.QTextEdit(self.tab_2)
        self.textEdit.setGeometry(QtCore.QRect(30, 10, 200, 61))
        self.textEdit.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: grey;\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.textEdit.setObjectName("textEdit")
        self.textEdit_2 = QtWidgets.QTextEdit(self.tab_2)
        self.textEdit_2.setGeometry(QtCore.QRect(250, 10, 200, 61))
        self.textEdit_2.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: grey;\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.textEdit_2.setObjectName("textEdit_2")
        self.textEdit_3 = QtWidgets.QTextEdit(self.tab_2)
        self.textEdit_3.setGeometry(QtCore.QRect(470, 10, 200, 61))
        self.textEdit_3.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: grey;\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.textEdit_3.setObjectName("textEdit_3")
        self.textEdit_4 = QtWidgets.QTextEdit(self.tab_2)
        self.textEdit_4.setGeometry(QtCore.QRect(690, 10, 200, 61))
        self.textEdit_4.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: grey;\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.textEdit_4.setObjectName("textEdit_4")
        self.textEdit_5 = QtWidgets.QTextEdit(self.tab_2)
        self.textEdit_5.setGeometry(QtCore.QRect(910, 10, 200, 61))
        font = QtGui.QFont()
        font.setBold(True)
        font.setItalic(False)
        font.setWeight(75)
        self.textEdit_5.setFont(font)
        self.textEdit_5.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: grey;\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.textEdit_5.setObjectName("textEdit_5")
        self.radioButton = QtWidgets.QRadioButton(self.tab_2)
        self.radioButton.setGeometry(QtCore.QRect(30, 90, 241, 61))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setWeight(75)
        self.radioButton.setFont(font)
        self.radioButton.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: rgb(151, 180, 255, 255);\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;\n"
"")
        self.radioButton.setObjectName("radioButton")
        self.radioButton_2 = QtWidgets.QRadioButton(self.tab_2)
        self.radioButton_2.setGeometry(QtCore.QRect(290, 90, 241, 61))
        font = QtGui.QFont()
        font.setPointSize(18)
        font.setBold(True)
        font.setWeight(75)
        self.radioButton_2.setFont(font)
        self.radioButton_2.setStyleSheet("background-color: rgb(255, 255, 255);\n"
"border-style: inset;\n"
"color: rgb(151, 180, 255, 255);\n"
"border-style: hidden;\n"
"border-width: 3px;\n"
"border-radius: 10px;\n"
"font: bold 14p;\n"
"min-width: 10em;\n"
"padding: 6px;")
        self.radioButton_2.setObjectName("radioButton_2")
        self.tabWidget.addTab(self.tab_2, "")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(0, -20, 1161, 681))
        self.label_2.setStyleSheet("background-color: qlineargradient(spread:pad, x1:1, y1:1, x2:1, y2:0, stop:0.00568182 rgba(255, 255, 255, 255), stop:1 rgba(151, 180, 255, 255));")
        self.label_2.setText("")
        self.label_2.setObjectName("label_2")
        self.label_2.raise_()
        self.tabWidget.raise_()
        MainWindow.setCentralWidget(self.centralwidget)


        self.retranslateUi(MainWindow)
        self.tabWidget.setCurrentIndex(0)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.mythread = MyThread()
        self.add_functions()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.pushButton.setText(_translate("MainWindow", "Добавить файл"))
        self.textBrowser.setPlaceholderText(_translate("MainWindow", "Путь к файлу"))
        self.pushButton_2.setText(_translate("MainWindow", "Загрузить"))
        self.zxc.setText(_translate("MainWindow", "Очистить"))
        self.pushButton_3.setText(_translate("MainWindow", "Удалить БД"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab), _translate("MainWindow", "Загрузка"))
        self.pushButton_5.setText(_translate("MainWindow", "Добавить"))
        self.textEdit.setPlaceholderText(_translate("MainWindow", "Фамалия"))
        self.textEdit_2.setPlaceholderText(_translate("MainWindow", "Имя"))
        self.textEdit_3.setPlaceholderText(_translate("MainWindow", "Отчество"))
        self.textEdit_4.setPlaceholderText(_translate("MainWindow", "Балл"))
        self.textEdit_5.setPlaceholderText(_translate("MainWindow", "Специальность"))
        self.radioButton.setText(_translate("MainWindow", "Бюджет"))
        self.radioButton_2.setText(_translate("MainWindow", "Коммерция"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_2), _translate("MainWindow", "Внесение"))


    def add_functions(self): # Функция для срабатывания кнопок
        self.pushButton.clicked.connect(self.openDirectory)
        self.pushButton_2.clicked.connect(self.on_clicked) 
        self.mythread.started.connect(self.on_clicked)
        self.mythread.finished.connect(self.on_finished)
        self.pushButton_3.clicked.connect(self.DLDB)
        self.zxc.clicked.connect(self.ClearText) 
        self.pushButton_5.clicked.connect(self.AddStudent)


    def openDirectory(self): # Функция открытия директории для выбора и дальнейшей работы с файлом
        try:
            file_name.clear()
            dir, _ = QFileDialog.getOpenFileName(None, "Выбрать файл", "~", "All Files (*);;sqlFiles (*.db);;exelFiles (*.xlsx);; wordFile (*.docx);; OldWord (*.doc)")
            url = QUrl.fromLocalFile(dir)
            filename = QFileInfo(dir).fileName()
            self.textBrowser.setText(f'{dir}')
            print(f'{dir}')
            file_name.append(f'{dir}')
            print(file_name)
        except FileNotFoundError:
            print("No such file")


    def on_change(self, s):
        self.label.setText(s)
    def on_clicked(self):
        self.pushButton_2.setDisabled(True)
        self.mythread.start()   
        self.mythread.mysignal.connect(self.on_change, QtCore.Qt.QueuedConnection)
    def on_finished(self):
        self.pushButton_2.setDisabled(False)
        file_name.clear()
        self.textBrowser.clear()
    def AddStudent(self): # Функция поштучного добавления студентов в БД
        import sqlite3
        try:
          db = sqlite3.connect('FIS.db')
          if self.radioButton.isChecked():
            cur = db.cursor()
            textFam = self.textEdit.toPlainText()
            textIm = self.textEdit_2.toPlainText()
            textOtch = self.textEdit_3.toPlainText()
            textBall = float(self.textEdit_4.toPlainText())
            textSpets = self.textEdit_5.toPlainText()

            cur.execute("INSERT INTO students (fam, imy, otch, ball, specialnost, forma) VALUES (?, ?, ?, ?, ?, ?)", (textFam, textIm, textOtch, textBall, textSpets, "Бюджет"))
            print("Работает")
            self.textEdit.clear()
            self.textEdit_2.clear()
            self.textEdit_3.clear()
            self.textEdit_4.clear()
            self.textEdit_5.clear()
            db.commit()
            db.close()
          if self.radioButton_2.isChecked():
            cur = db.cursor()
            textFam = self.textEdit.toPlainText()
            textIm = self.textEdit_2.toPlainText()
            textOtch = self.textEdit_3.toPlainText()
            textBall = float(self.textEdit_4.toPlainText())
            textSpets = self.textEdit_5.toPlainText()

            cur.execute("INSERT INTO students (fam, imy, otch, ball, specialnost, forma) VALUES (?, ?, ?, ?, ?, ?)", (textFam, textIm, textOtch, textBall, textSpets, "Коммерция"))
            print("Тоже работает")
            self.textEdit.clear()
            self.textEdit_2.clear()
            self.textEdit_3.clear()
            self.textEdit_4.clear()
            self.textEdit_5.clear()

            msgBase = QMessageBox()
            msgBase.setWindowTitle("Информация")
            msgBase.setText("Данные были успешно внесены!")
            msgBase.setIcon(QMessageBox.Information)
            msgBase.exec_()

            db.commit()
            db.close()
        except ValueError:
                  msg8 = QMessageBox()
                  msg8.setWindowTitle("Информация")
                  msg8.setText("Какие-то из данных были введены неверно. Проверьте пожалуйста")
                  msg8.setIcon(QMessageBox.Critical)
                  msg8.exec_()
    def DLDB (self): # Функция удаления данных из БД
        msg = QMessageBox.warning(None, 'Warning', "Вы хотите удалить Базу Данных.\nВы уверены?", QMessageBox.Ok | QMessageBox.Cancel)
        if msg == QMessageBox.Ok:
            import sqlite3
            db = sqlite3.connect('FIS.db')
            cur = db.cursor()
            cur.execute("DELETE FROM students")
            db.commit()
            db.close()
            msg2 = QMessageBox()
            msg2.setWindowTitle("Information")
            msg2.setText("База данных была успешно удалена!")
            msg2.setIcon(QMessageBox.Information)
            msg2.setStandardButtons(QMessageBox.Ok)
            x2 = msg2.exec_()
        elif msg == QMessageBox.Cancel:
            pass

    def ClearText(self): # Очищение текста в окне и массива
      try:
        file_name.clear()
        self.textBrowser.clear()
      except IndexError:
        pass


if __name__ == "__main__": # Отрисовка/показ окон/интерфейса
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())

    
def parsing():
    headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:105.0) Gecko/20100101 Firefox/105.0'
}
    def url(page):
        url = f"http://prof.mo.mosreg.ru/api/spoPetition/search/advancedSearch?page={page}&size=500&sort=createdTs%2Cdesc&order=desc&q=%7B%22spoEducationYear%22%3A%224%22%2C%22status%22%3A%22ACCEPTED%22%7D&projection=grid"
        return(url)
    # {username: "korshikova", password: "4614edba3fdb0b76554a2ed52af09b2c26d721f64caebc9d39cbc33e93e215b3"}



    import requests
    import json
    import requests
    from bs4 import BeautifulSoup as bs
    import pandas as pd
    import requests
    import time
    from bs4 import BeautifulSoup

    def login():
        url = 'http://prof.mo.mosreg.ru/api/login'
        s = requests.Session()
        payload = {
            'username': "korshikova",
            'password': "4614edba3fdb0b76554a2ed52af09b2c26d721f64caebc9d39cbc33e93e215b3",
        }
        res = s.post('http://prof.mo.mosreg.ru/api/login', json=payload,headers=headers)

        s.headers.update(json.loads(res.content))

        token_new = str(s.headers['token'])
        with open(f"token.txt", "w", encoding='utf-8') as file:
                file.write(token_new)
        with open("token.txt", "r") as f:
                token_old = f.read()


    def pars_data():
        with open("token.txt", "r") as f:
            token_old = f.read()
        s = requests.Session()
        headers2 = {
        'user-agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:101.0) Gecko/20100101 Firefox/101.0',
        'Content-Type': 'application/json; charset=UTF-8',
        'Connection': 'Keep-Alive',
        'Cookie': f'auth-token=Token%20{token_old}',
        }
        response = s.get(url(page=0),headers=headers2)
        if response.status_code != 401:
            for page in range(response.status_code):
                url_pars = url(page)
                # cookies = "Cookie: auth-token=Token%20bmhMd0M0VGhFMzhIWnpWNGxoNTUwMGQybVRLUGY0VUxERGRsRlg2UHlXN090UUVDU3JnMGlHK2dBbGgxTzRZMDVLTDk2V3llTzVaNjI2cTRVYTU4aWc9PQ; _ym_visorc=w; _ym_isad=2"
                response = s.get(url_pars,headers=headers2) 
                soup = BeautifulSoup(response.text, 'lxml')
                with open("test.json", "w", encoding='utf-8') as file:
                        file.write(soup.text + '\n')
                with open('test.json', encoding='utf-8') as f:
                        templates = json.load(f)
                my_json_str = json.dumps(templates)
                json_var = json.loads(my_json_str)
                souplen = len(json_var['_embedded']['spoPetitions'])
                # В одном full файле должно быть 33026-33029 строк 
                if souplen > 1:
                    with open(f"list{page}.json", "w", encoding='utf-8') as file:
                        file.write(soup.text + '\n')
                else:
                    break
        else:
            login()
            pars_data()



    pars_data()


def db_add1():
    import sqlite3
    import requests
    import json
    import requests
    from bs4 import BeautifulSoup as bs
    import pandas as pd
    import requests
    import time
    from bs4 import BeautifulSoup
    db = sqlite3.connect('FIS.db')
    cur = db.cursor()
    def dobavit_table4(id,imy,fam,Otchestvo,gender,snils,email,tspodachi,needDormitory,dr,drmesto,idmestodr,series,number,school,ball,diplonend,idcountre,iddocumenttype,idpasport,idspec,datapodachi,status,idzayv,idman,kemvidan,kogdavidan,idatestat,specname,forma):
        cur.execute('INSERT INTO table1 (id,imy,fam,Otchestvo,gender,snils,email,tspodachi,needDormitory,dr,drmesto,idmesto,series,number,school,ball,diplomend,idcountry,iddicomenttype,idpasport,idspec,datapodachi,status,idzayv,idman,kemvidan,kogdavidan,idatestat,specname,forma) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',[id,imy,fam,Otchestvo,gender,snils,email,tspodachi,needDormitory,dr,drmesto,idmestodr,series,number,school,ball,diplonend,idcountre,iddocumenttype,idpasport,idspec,datapodachi,status,idzayv,idman,kemvidan,kogdavidan,idatestat,specname,forma]) 

        db.commit()

    text_all_txt = []
    xml_txt = []
    with open("data.txt", "r", encoding='utf-8') as file:
        line=file.read().split("\n")
        for i in range(len(line)-1):
            url = line[i]


            with open("token.txt", "r") as f:
                token_old = f.read()
            s = requests.Session()
            headers2 = {
            'user-agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:101.0) Gecko/20100101 Firefox/101.0',
            'Content-Type': 'application/json; charset=UTF-8',
            'Connection': 'Keep-Alive',
            'Cookie': f'auth-token=Token%20{token_old}',
            }
            response = s.get(url,headers=headers2)
            if response.status_code != 401:
                    url_pars = url
                    response = s.get(url_pars,headers=headers2) 
                    soup = BeautifulSoup(response.text, 'lxml')
                    with open("test.json", "w", encoding='utf-8') as file:
                            file.write(soup.text)


                    with open('test.json', encoding='utf-8') as f:
                        templates = json.load(f)
                        text = templates
                        my_json_str = json.dumps(text)
                        json_var = json.loads(my_json_str)
                        # код 
                        # kod
                        fio = ["firstName","lastName","middleName"]
                        id = str(json_var['id'])
                        datafio = []
                        idzayv = str(f'20221308{i}')
                        for i in range(len(fio)):
                            FIO = (json_var['document'][fio[i]])
                            datafio.append(FIO)  
                        gender = str(json_var['gender'])
                        snils = str(json_var['snils'])
                        email = str(json_var['email'])
                        tspodachi = str(json_var['createdTs'])
                        needDormitory = str(json_var['needDormitory'])
                        dr = str(json_var['document']['birthDate'])
                        drmesto = str(json_var['document']['birthPlace'])
                        idmestodr = str(json_var['document']['nationality']['code'])
                        series = str(json_var['document']['series'])
                        number = str(json_var['document']['number'])
                        school = str(json_var['basicEducationOrganizationPlain'])
                        ball = str(json_var['educationDocumentGPA'])
                        diplonend = str(json_var['basicEducationEndYear'])
                        idcountre = str(json_var['document']['nationality']['code'])
                        iddocumenttype = str(json_var['document']['type']['category']['documentType']['code'])
                        idpasport = str(json_var['document']['id'])
                        idspec = str(json_var['admissionPlan']['id'])
                        datapodachi = str(json_var['createdTs'])
                        status = str(json_var['educationDocument']['type']['category']['code'])
                        id = str(json_var['id'])
                        idman = str(json_var['document']['id'])
                        kemvidan = str(json_var['document']['source'])
                        kogdavidan = str(json_var['document']['dateObtain'])
                        idatestat = (json_var['attaches'])
                        my_json_str2 = json.dumps(idatestat)
                        json_var2 = json.loads(my_json_str2)
                        json_var2 = json_var2[1 : -1]
                        json_var2 = (dict[json_var2])
                        idatestat = str(json_var['id'])
                        fam = str(datafio[1])
                        imy = str(datafio[0])
                        specname = str(json_var['admissionPlan']['specialityComputedName'])
                        Otchestvo = str(datafio[2])
                        forma = str(json_var['financialType']['title'])
                        dobavit_table4(id,imy,fam,Otchestvo,gender,snils,email,tspodachi,needDormitory,dr,drmesto,idmestodr,series,number,school,ball,diplonend,idcountre,iddocumenttype,idpasport,idspec,datapodachi,'4',idzayv,idman,kemvidan,kogdavidan,idatestat,specname,forma)

def db_add2Office(): # таблицы
    import os
    import sqlite3
    import openpyxl



    db = sqlite3.connect('FIS.db')
    cur = db.cursor()
  # заменить имя файла
    file_to_read = openpyxl.load_workbook(data_only=True)
    sheet = file_to_read['Лист1']

    if str(file_to_read.sheetnames) == "['Лист1']":
      for row in range(2, sheet.max_row + 1):
        data = []
        for col in range(1, 8):
            value = sheet.cell(row, col).value
            data.append(value)
        ball = str(data[4])
        if len(ball) < 2:
              ball = ball + '.0'
        cur.execute("INSERT INTO students (fam, imy, otch, ball, specialnost, forma) VALUES (?, ?, ?, ?, ?, ?)", (data[1], data[2], data[3], ball, data[5], data[6]))

      db.commit()
      db.close()
    else:
      print("Некорректное название листа")


def db_adddocx(choses_file):
  import sqlite3
  import docx

  db = sqlite3.connect('FIS.db')
  cur = db.cursor()
  # file_name = r'C:\Users\serge\Desktop\students.docx'
  file_name = choses_file
  doc = docx.Document(file_name)
  text = []

  for table in doc.tables:
    for row in table.rows:
      x = []
      if len(x) < 5:
        for cell in row.cells:
# cell.text
          x.append(cell.text)
      else:
        cur.execute("INSERT INTO students (fam, imy, otch, ball, specialnost, forma) VALUES (?, ?, ?, ?, ?, ?)", (x[1], x[2], x[3], x[4], x[5], x[6]))
        db.commit()
  db.close()


def xml():
  asd = []
  # имя + фам + отч + балл == 0.2%
  # имя + фам + отч + балл + tspodachi + spec 
  import time
  lallal = []
  from itertools import groupby
  dubl= []
  origsnils =[]
  import sqlite3
  db = sqlite3.connect('FIS.db')
  cur = db.cursor()
  def snils1():

      # Снилсы которые появляются больше 1 раза
      cur.execute('SELECT * FROM table1')

      result = cur.fetchall()
      for i in range(len(result)):
        dubl.append(result[i])
  snils1()
  old_x =[]
  def snils2():
        
        nex_sp = []
        for i in range(len(dubl)):
          if len(dubl[i][5]) > 0:
            cur.execute(f'SELECT imy,fam,Otchestvo,ball,specname,forma,tspodachi FROM table1 where imy = "{dubl[i][1]}" and fam = "{dubl[i][2]}" and Otchestvo = "{dubl[i][3]}" AND ball = "{dubl[i][15]}"')
            result = cur.fetchall()
            old_x.append(result)
        new_x = [el for el, _ in groupby(old_x)]
        for i in range(len(new_x)):
          for z in range(len(new_x[i])):
            cur.execute(f'SELECT imy,fam,otch,ball,specialnost,forma FROM students where imy = "{new_x[i][z][0]}" and fam = "{new_x[i][z][1]}" and otch = "{new_x[i][z][2]}" AND ball = "{new_x[i][z][3]}"')
            resultSP = cur.fetchall()
            if len((resultSP)) >0:
              nex_sp.append(resultSP)
        new_sp = [el for el, _ in groupby(nex_sp)]
        for i in range(len(new_sp)):
          origsnils.append(new_sp[i])


  snils2()
  # print(len(old_x))

  cur.execute(f'SELECT * FROM table1')
  xmldata = cur.fetchall()
  # for i in range(len(xmldata)):

  #     # print(xmldata[i][28])
  #       if xmldata[i][29] == 'Бюджет':
  #           s = str(xmldata[i][28])
  #           s1 = "".join(c for c in s if  c.isdecimal())
  #           idspec = (s1+'1'+'22')
  #       else:
  #           s = str(xmldata[i][28])
  #           s1 = "".join(c for c in s if  c.isdecimal())
  #           idspec = (s1+'2'+'22')
  aaa = []
  for i in range(len(old_x)):
    aaa.append(old_x[i])


  new_x2 = [el for el, _ in groupby(aaa)]

  def join_strs(strs):
      result = ''
      for s in strs:
          result += '\n            ' + s  # а вот и плюс
      return result[1:]

  # если среди new_x2 есть кто то из оригинала то он основной 

  for i in range(len(new_x2)):
    cxz = []
    for z in range(len(new_x2[i])):
        cur.execute(f'SELECT imy,fam,otch,ball,specialnost,forma FROM students where imy = "{new_x2[i][z][0]}" and fam = "{new_x2[i][z][1]}" and otch = "{new_x2[i][z][2]}" AND ball = "{new_x2[i][z][3]}" and specialnost = "{new_x2[i][z][4]}" and forma = "{new_x2[i][z][5]}"')
        resultSP = cur.fetchall()
        
        if len((resultSP)) > 0:
                if new_x2[i][z][5] == 'Бюджет':
                    s = str(new_x2[i][z][4])
                    s1 = "".join(c for c in s if  c.isdecimal())
                    idspec = (s1+'1'+'22')
                    full = [idspec,new_x2[i][z][6]]
                    cxz.append(full)



                else:
                    s = str(new_x2[i][z][4])
                    s1 = "".join(c for c in s if  c.isdecimal())
                    idspec = (s1+'2'+'22')
                    full = idspec
                    cxz.append(full)
        elif len((resultSP)) < 1:
              if new_x2[i][z][5] == 'Бюджет':
                  s = str(new_x2[i][z][4])
                  s1 = "".join(c for c in s if  c.isdecimal())
                  idspec = (s1+'1'+'22')
                  full = idspec
                  cxz.append(full)

              else:
                  s = str(new_x2[i][z][4])
                  s1 = "".join(c for c in s if  c.isdecimal())
                  idspec = (s1+'2'+'22')
                  full = idspec
                  cxz.append(full)

    zc = []
    cxz = [cxz]

    for i in range(len(cxz)):
      for z in range(len(cxz[i])):
        if len(cxz[i][z]) >8:
          one = f'<FinSourceEduForm><CompetitiveGroupUID>{cxz[i][z]}</CompetitiveGroupUID></FinSourceEduForm>'
          zc.append(str(one))
        elif len(cxz[i][z]) < 3:
          two = f'<FinSourceEduForm><CompetitiveGroupUID>{cxz[i][z][0]}</CompetitiveGroupUID><IsAgreedDate>{cxz[i][z][1]}</IsAgreedDate></FinSourceEduForm>'
          zc.append(two)
    asd.append(zc)
  for i in range(len(asd)):
    cur.execute(f'SELECT * FROM table1 where imy = "{new_x2[i][0][0]}" and fam = "{new_x2[i][0][1]}" and Otchestvo = "{new_x2[i][0][2]}" AND ball = "{new_x2[i][0][3]}" and specname = "{new_x2[i][0][4]}" and forma = "{new_x2[i][0][5]}"')
    resultSP2 = cur.fetchall()
    
    xml = f"""        
          <Application>
            <UID>{resultSP2[0][0]}</UID>
            <ApplicationNumber>{resultSP2[0][23]}</ApplicationNumber>
            <Entrant>
              <UID>{resultSP2[0][23]}</UID>
              <LastName>{resultSP2[0][2]}</LastName>
              <FirstName>{resultSP2[0][1]}</FirstName>
              <MiddleName>{resultSP2[0][3]}</MiddleName>
              <GenderID>{resultSP2[0][4]}</GenderID>
              <SNILS>{resultSP2[0][5]}</SNILS>
              <EmailOrMailAddress>
                <Email>{resultSP2[0][6]}</Email>
              </EmailOrMailAddress>
            </Entrant>
            <RegistrationDate>{resultSP2[0][7]}</RegistrationDate>
            <NeedHostel>{resultSP2[0][8]}</NeedHostel>
            <StatusID>{resultSP2[0][21]}</StatusID>
            <FinSourceAndEduForms>
            <FinSourceEduForm>
  {join_strs(asd[i])}
            </FinSourceEduForm>
            </FinSourceAndEduForms>
            <ApplicationDocuments>
              <IdentityDocument>
                <UID>{resultSP2[0][19]}</UID>
                <DocumentSeries>{resultSP2[0][12]}</DocumentSeries>
                <DocumentNumber>{resultSP2[0][13]}</DocumentNumber>
                <DocumentDate>{resultSP2[0][25]}</DocumentDate>
                <DocumentOrganization>{resultSP2[0][24]}</DocumentOrganization>
                <IdentityDocumentTypeID>{resultSP2[0][18]}</IdentityDocumentTypeID>
            <NationalityTypeID>{resultSP2[0][17]}</NationalityTypeID>
                <BirthDate>{resultSP2[0][9]}</BirthDate>
                <ReleasePlace>{resultSP2[0][10]}</ReleasePlace>
            <ReleaseCountryID>{resultSP2[0][11]}</ReleaseCountryID>
              </IdentityDocument>
              <EduDocuments>
                <EduDocument>
                  <SchoolCertificateBasicDocument>
                    <UID>{resultSP2[0][26]}</UID>
                    <DocumentSeries>{resultSP2[0][12]}</DocumentSeries>
                    <DocumentNumber>{resultSP2[0][13]}</DocumentNumber>
                    <DocumentDate>{resultSP2[0][16]}</DocumentDate>
                    <DocumentOrganization>{resultSP2[0][14]}</DocumentOrganization>
                    <GPA>{resultSP2[0][15]}</GPA>
                  </SchoolCertificateBasicDocument>
                </EduDocument>
              </EduDocuments>
            </ApplicationDocuments>
          </Application>
  """   
    # print(xml)
    # time.sleep(15)
    with open("xmlAll.txt", "a", encoding='utf-8') as file:
            file.write(xml)


def data():
    import json

    i = 0 
    for i in range(4):
        x = []
        with open(f'list{i}.json', encoding='utf-8') as f:
            templates = json.load(f)
        text = templates

        my_json_str = json.dumps(text)
        json_var = json.loads(my_json_str)

        text = (json_var['_embedded']['spoPetitions'])
        # text = str(text)
        for i in range(len(text)):
            # print(text[i])
            txt = text[i]
            data = (txt["_links"]['self']['href'])
            #   print(data + '?projection=detail')
            #   print(i)
            x.append(data)


        with open("data.txt", "a", encoding='utf-8') as file:
                    for i in range(len(x)):
                        file.write(str(x[i]+ '?projection=detail') + '\n')


# deleteBd()
# db_add2Office()
# db_adddocx()
